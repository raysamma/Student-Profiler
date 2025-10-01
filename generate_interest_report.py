import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import os

# Load the Excel file
excel_path = "Final data for the script.xlsx"  # Ensure this file is in the same directory
xls = pd.ExcelFile(excel_path)

# Load sheets
intro_text = xls.parse("Introduction").iloc[:, 0].dropna().tolist()
scores_df = xls.parse("Scores")
percentile_table = xls.parse("Percentile Table")
jobs_df = xls.parse("RIASEC Codes")
# Clean up the 'Code' column to remove any extra spaces that could cause matching errors.
jobs_df['Code'] = jobs_df['Code'].str.strip()


# Load personality descriptions
descriptions = {}
for trait in ["Realistic", "Investigative", "Artistic", "Social", "Enterprising", "Conventional"]:
    df = xls.parse(trait)
    
    # FIX: For ALL sheets, filter out rows that only contain numbers.
    # This prevents stray numbers from appearing in the descriptions.
    is_text_mask = pd.to_numeric(df.iloc[:, 0], errors='coerce').isna()
    descriptions[trait[0]] = df[is_text_mask].iloc[:, 0].dropna().tolist()


# Convert percentile table into a lookup dictionary
percentile_lookup = {}
for index, row in percentile_table.iterrows():
    try:
        score = int(row.iloc[0])
        percentiles = row[1:].to_dict()
        percentile_lookup[score] = percentiles
    except (ValueError, TypeError):
        continue

# Trait initials to full names
trait_map = {
    'R': 'Realistic',
    'I': 'Investigative',
    'A': 'Artistic',
    'S': 'Social',
    'E': 'Enterprising',
    'C': 'Conventional'
}

# Convert raw scores to percentiles
def get_percentiles(raw_scores):
    percentiles = {}
    for trait, score in raw_scores.items():
        percentile = percentile_lookup.get(score, {}).get(trait, 0)
        percentiles[trait] = percentile
    return percentiles

# Plot graph for each student
def plot_percentiles(percentiles, student_id):
    traits = ['R', 'I', 'A', 'S', 'E', 'C']
    values = [percentiles.get(t, 0) for t in traits]
    plt.figure(figsize=(6, 3))
    plt.bar(traits, values, color='skyblue')
    plt.ylim(0, 100)
    plt.title('RIASEC Percentile Scores')
    plt.xlabel('Personality Type')
    plt.ylabel('Percentile')
    graph_path = f"graph_{student_id}.png"
    plt.tight_layout()
    plt.savefig(graph_path)
    plt.close()
    return graph_path

# Create the Word document
doc = Document()

# Generate reports for each student
for idx, row in scores_df.iterrows():
    name = row['Student Name']
    roll = row['Roll Number']
    clas = row['Class']
    section = row.get(' Section ', '') # Use .get() for safety

    # FIX: Correctly read all six scores from the row.
    raw_scores = {
        'R': row['Realistic'],
        'I': row['Investigative'],
        'A': row['Artistic'],
        'S': row['Social'],
        'E': row['Enterprising'],
        'C': row['Conventional']
    }

    percentiles = get_percentiles(raw_scores)
    sorted_types = sorted(percentiles.items(), key=lambda x: x[1], reverse=True)
    top3 = sorted_types[:3]
    # CORRECT LOGIC: Use the code based on the strength of the scores. No more alphabetizing.
    theme_code = ''.join([x[0] for x in top3])


    # Cover Page
    if idx > 0:
        doc.add_page_break()
    doc.add_heading("Holland Career Interest Assessment Report", 0)
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Class: {clas}")
    doc.add_paragraph(f"Section: {section}")
    doc.add_paragraph(f"Roll Number: {roll}")

    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    for para in intro_text:
        doc.add_paragraph(str(para))

    # 2. Score Analysis
    doc.add_heading("2. Score Analysis", level=1)
    doc.add_paragraph(f"Your Holland Code: {theme_code}")
    
    score_table = doc.add_table(rows=1, cols=3)
    score_table.style = 'Table Grid'
    hdr_cells = score_table.rows[0].cells
    hdr_cells[0].text = 'Type'
    hdr_cells[1].text = 'Raw Score'
    hdr_cells[2].text = 'Percentile'
    for t_initial in ['R', 'I', 'A', 'S', 'E', 'C']:
        row_cells = score_table.add_row().cells
        row_cells[0].text = trait_map[t_initial]
        row_cells[1].text = str(raw_scores[t_initial])
        row_cells[2].text = str(percentiles.get(t_initial, 0))

    # Graph
    graph_path = plot_percentiles(percentiles, roll)
    doc.add_picture(graph_path, width=Inches(6))
    os.remove(graph_path)

    # 3. Top 3 Personality Descriptions
    doc.add_heading("3. Your Top 3 Personality Themes (In Order of Strength)", level=1)
    for i, (t_initial, percentile_val) in enumerate(top3):
        full_name = trait_map.get(t_initial, "Unknown")
        # Add the heading and control its spacing
        heading = doc.add_heading(f"{full_name} ({t_initial})", level=2)
        heading.paragraph_format.space_after = Pt(6) # Add a small space after the heading
        
        # --- ADD THEME IMAGE FROM 'images' FOLDER ---
        image_path = os.path.join('images', f'{full_name}.png')
        if os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(6))
            # Control the spacing of the paragraph containing the picture
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.paragraph_format.space_before = Pt(6)
            last_paragraph.paragraph_format.space_after = Pt(12)
        else:
            print(f"Warning: Image not found at {image_path}")

        # The rest of the lines are the description paragraphs
        for line in descriptions.get(t_initial, [])[1:]:
            doc.add_paragraph(str(line))
        
        # Add a page break after the first and second themes, but not the last.
        if i < len(top3) - 1:
            doc.add_page_break()

    # 4. Career Recommendations (FINAL HYBRID LOGIC)
    doc.add_heading("4. Career Recommendations", level=1)
    # Find the specific row for the student's theme code
    code_info = jobs_df[jobs_df['Code'] == theme_code]
    
    if code_info.empty:
        # If no exact 3-letter match, provide a more helpful message
        strongest_theme_initial = top3[0][0]
        strongest_theme_full = trait_map.get(strongest_theme_initial)
        doc.add_paragraph(f"No specific job recommendations were found for your unique code ({theme_code}).")
        doc.add_paragraph(f"This is a great opportunity to explore careers that align with your strongest personality theme: {strongest_theme_full} ({strongest_theme_initial}). Consider researching jobs related to this primary interest area.")
    else:
        # --- LOGIC FOR WORK FIELDS (As per your instructions) ---
        target_code_id = code_info.iloc[0, 0]
        work_field_matches = jobs_df[jobs_df.iloc[:, 4] == target_code_id]
        recommended_work_fields = work_field_matches['Work Field'].dropna().unique().tolist()

        # --- HYBRID LOGIC FOR JOBS (As per your instructions) ---
        # Attempt 1: Use the work_id linking method
        target_work_id = code_info.iloc[0, 3]
        job_matches = jobs_df[jobs_df.iloc[:, 7] == target_work_id]
        
        # Fallback: If Attempt 1 finds no jobs, use the direct code match
        if job_matches.empty:
            job_matches = code_info

        # --- Displaying the results in a grouped format ---
        doc.add_paragraph(f"Based on your Holland Code ({theme_code}), here are some career paths to explore:")
        p = doc.add_paragraph()
        p.add_run("Note: The Job Zone indicates the level of preparation needed for a job (1 = little preparation, 5 = extensive preparation).").italic = True

        # Group the matched jobs by their "Work Field"
        grouped_jobs = job_matches.groupby('Work Field')

        for work_field, jobs_in_field in grouped_jobs:
            # Add the work field as a sub-heading
            p = doc.add_paragraph()
            p.add_run(work_field.strip().replace('\n', ' ')).bold = True
            
            # List the jobs under that work field
            for index, job_row in jobs_in_field.iterrows():
                job = job_row['Jobs']
                # Note the trailing space in 'job zones ' which is present in the CSV header
                zone = job_row['job zones ']
                if pd.notna(job):
                    try:
                        # Attempt to format the zone as a whole number
                        zone_str = f"{int(float(zone))}"
                    except (ValueError, TypeError):
                        # If it's not a number, use it as is
                        zone_str = str(zone).strip()
                    doc.add_paragraph(f"• {job.strip()} (Job Zone: {zone_str})")

    # 5. Skill Enhancement Suggestions
    doc.add_heading("5. How to Explore Your Interests", level=1)
    doc.add_paragraph("• Participate in school clubs, after-school activities, or personal projects that align with your top traits.")
    doc.add_paragraph("• Talk to people! Seek guidance from teachers, counselors, or professionals working in fields that interest you.")
    doc.add_paragraph("• Volunteer or find a part-time job in an area related to your interests to get real-world experience.")

    # 6. Interest Report Summary
    doc.add_heading("6. Interest Report Summary", level=1)
    top_themes_full = [trait_map.get(t[0]) for t in top3]
    doc.add_paragraph(f"Based on your top scores in {', '.join(top_themes_full)} ({theme_code}), your interests align with traits and work environments typical for these personality types. Use this report as a starting point to explore, learn, and reflect on what truly motivates you.")

    # 7. Closing Note
    doc.add_heading("7. A Final Note", level=1)
    doc.add_paragraph("Understanding your interests is a powerful first step toward discovering a fulfilling career path. This report is a snapshot in time. Continue to be curious, keep learning, and trust in your unique potential. Your interests may evolve, and that's a natural part of growth!")

# Save final compiled report
output_path = "Holland_Interest_Assessment_Report_Generated.docx"
doc.save(output_path)
print(f"\n✅ Report generated: {output_path}")
